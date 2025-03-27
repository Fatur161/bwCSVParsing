import sys
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

if len(sys.argv) != 3:
    print("Использование: python script.py путь_к_csv название_выходного_файла.docx")
    sys.exit(1)

# Получаем пути из аргументов командной строки
csv_file_path = sys.argv[1]
docx_file_path = sys.argv[2] + ".docx"

# Проверка расширения файлов
if not csv_file_path.endswith('.csv'):
    print("Входной файл должен быть формата .csv")
    sys.exit(1)

# Чтение CSV файла
try:
    df = pd.read_csv(csv_file_path)
except FileNotFoundError:
    print(f"Файл {csv_file_path} не найден.")
    sys.exit(1)
except pd.errors.EmptyDataError:
    print("Файл пуст.")
    sys.exit(1)

# Создание нового DOCX документа
doc = Document()

# Настройка стилей
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
style.paragraph_format.space_after = Pt(6)

# Добавление заголовка
heading = doc.add_heading('Data Экспорт', level=1)
heading.style.font.name = 'Times New Roman'
heading.style.font.size = Pt(14)

# Проход по строкам DataFrame и добавление данных в DOCX
for index, row in df.iterrows():
    current_date = datetime.now().strftime("%d.%m.%Y")
    section_heading = doc.add_heading(f'Запись {index + 1} от {current_date}', level=2)
    section_heading.style.font.name = 'Times New Roman'
    section_heading.style.font.size = Pt(13)

    # Добавляем каждое поле с новой строки
    for column in df.columns:
        value = row.get(column, 'N/A')
        paragraph = doc.add_paragraph(f'{column}: {value}')
        paragraph.style = style

    # Добавляем разделитель
    doc.add_paragraph('-' * 25)

# Сохранение DOCX файла
doc.save(docx_file_path)
print(f"Данные успешно экспортированы в {docx_file_path}")
